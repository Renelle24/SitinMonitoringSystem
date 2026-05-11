from flask import Flask, render_template, request, redirect, url_for, session, flash, jsonify, make_response
from flask_mysqldb import MySQL
from datetime import datetime, timedelta
import json
import csv
import io

app = Flask(__name__)
app.secret_key = 'ccs-sitin-secret-key-2026'

# ─── MySQL Configuration ───────────────────────────────────────────────────────
app.config['MYSQL_HOST'] = 'localhost'
app.config['MYSQL_USER'] = 'root'
app.config['MYSQL_PASSWORD'] = ''
app.config['MYSQL_DB'] = 'sitin_db'
app.config['MYSQL_PORT'] = 3306
app.config['MYSQL_CURSORCLASS'] = 'DictCursor'

mysql = MySQL(app)

# ─── Helpers ───────────────────────────────────────────────────────────────────

def find_student(student_id):
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM students WHERE id = %s", (student_id,))
    student = cur.fetchone()
    cur.close()
    return student

def get_announcements():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM announcements ORDER BY id DESC")
    rows = cur.fetchall()
    cur.close()
    return rows

def next_sit_id():
    cur = mysql.connection.cursor()
    cur.execute("SELECT COUNT(*) as cnt FROM sitin_records")
    row = cur.fetchone()
    cur.close()
    return f"SIT-{(row['cnt'] + 1):04d}"

# ─── Admin decorator ───────────────────────────────────────────────────────────

def admin_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get('role') != 'admin':
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

# ─── Student decorator ─────────────────────────────────────────────────────────

def student_required(f):
    from functools import wraps
    @wraps(f)
    def decorated(*args, **kwargs):
        if session.get('role') != 'student':
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

# ─── Auth routes ───────────────────────────────────────────────────────────────

@app.route('/')
def index():
    return redirect(url_for('login'))

@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        sid = request.form.get('id_number', '').strip()
        pwd = request.form.get('password', '').strip()

        if sid == 'admin' and pwd == 'admin':
            session['user'] = 'admin'
            session['role'] = 'admin'
            return redirect(url_for('admin_dashboard'))

        student = find_student(sid)
        if student and student['password'] == pwd:
            session['user'] = sid
            session['role'] = 'student'
            return redirect(url_for('student_dashboard'))

        flash('Invalid ID number or password.', 'danger')
    return render_template('login.html')

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        sid    = request.form.get('id_number', '').strip()
        last   = request.form.get('last_name', '').strip()
        first  = request.form.get('first_name', '').strip()
        middle = request.form.get('middle_name', '').strip()
        year   = int(request.form.get('year_level', 1))
        course = request.form.get('course', 'BSIT').strip()
        email  = request.form.get('email', '').strip()
        addr   = request.form.get('address', '').strip()
        pwd    = request.form.get('password', '').strip()
        pwd2   = request.form.get('repeat_password', '').strip()

        if find_student(sid):
            flash('ID Number already registered.', 'danger')
        elif pwd != pwd2:
            flash('Passwords do not match.', 'danger')
        elif not sid or not last or not first or not pwd:
            flash('Please fill all required fields.', 'danger')
        else:
            cur = mysql.connection.cursor()
            cur.execute("INSERT INTO students (id, last, first, middle, year, course, email, address, session, password) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                (sid, last, first, middle, year, course, email, addr, 30, pwd))
            mysql.connection.commit()
            cur.close()
            flash('Registration successful! Please login.', 'success')
            return redirect(url_for('login'))
    return render_template('register.html')

# ─── Admin routes ──────────────────────────────────────────────────────────────

@app.route('/admin')
@admin_required
def admin_dashboard():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM students")
    students = cur.fetchall()
    cur.execute("SELECT * FROM sitin_records")
    sitin_records = cur.fetchall()
    cur.close()

    currently = [r for r in sitin_records if r['status'] == 'Active']
    stats = {}
    for r in sitin_records:
        stats[r['purpose']] = stats.get(r['purpose'], 0) + 1

    return render_template('admin_dashboard.html',
        students=students,
        sitin_records=sitin_records,
        announcements=get_announcements(),
        currently_sitin=len(currently),
        total_sitin=len(sitin_records),
        stats=json.dumps(stats)
    )

@app.route('/admin/announce', methods=['POST'])
@admin_required
def post_announcement():
    body = request.form.get('body', '').strip()
    if body:
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO announcements (author, date, body) VALUES (%s,%s,%s)",
            ("CCS Admin", datetime.now().strftime("%Y-%b-%d"), body))
        mysql.connection.commit()
        cur.close()
    return redirect(url_for('admin_dashboard'))

@app.route('/admin/students')
@admin_required
def student_list():
    q = request.args.get('q', '').lower()
    cur = mysql.connection.cursor()
    if q:
        cur.execute("SELECT * FROM students WHERE LOWER(id) LIKE %s OR LOWER(CONCAT(first,' ',last)) LIKE %s",
            (f'%{q}%', f'%{q}%'))
    else:
        cur.execute("SELECT * FROM students")
    students = cur.fetchall()
    cur.close()
    return render_template('student_list.html', students=students, query=q, announcements=get_announcements())

@app.route('/admin/students/add', methods=['GET', 'POST'])
@admin_required
def add_student():
    if request.method == 'POST':
        sid = request.form.get('id_number', '').strip()
        if find_student(sid):
            flash('ID already exists.', 'danger')
        else:
            cur = mysql.connection.cursor()
            cur.execute("INSERT INTO students (id, last, first, middle, year, course, email, address, session, password) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                (sid,
                 request.form.get('last_name','').strip(),
                 request.form.get('first_name','').strip(),
                 request.form.get('middle_name','').strip(),
                 int(request.form.get('year_level', 1)),
                 request.form.get('course','BSIT').strip(),
                 request.form.get('email','').strip(),
                 request.form.get('address','').strip(),
                 int(request.form.get('session', 30)),
                 request.form.get('password','pass123').strip()))
            mysql.connection.commit()
            cur.close()
            flash('Student added.', 'success')
            return redirect(url_for('student_list'))
    return render_template('student_form.html', student=None, announcements=get_announcements())

@app.route('/admin/students/edit/<sid>', methods=['GET', 'POST'])
@admin_required
def edit_student(sid):
    student = find_student(sid)
    if not student:
        flash('Student not found.', 'danger')
        return redirect(url_for('student_list'))
    if request.method == 'POST':
        cur = mysql.connection.cursor()
        cur.execute("UPDATE students SET last=%s, first=%s, middle=%s, year=%s, course=%s, email=%s, address=%s, session=%s WHERE id=%s",
            (request.form.get('last_name','').strip(),
             request.form.get('first_name','').strip(),
             request.form.get('middle_name','').strip(),
             int(request.form.get('year_level', 1)),
             request.form.get('course','BSIT').strip(),
             request.form.get('email','').strip(),
             request.form.get('address','').strip(),
             int(request.form.get('session', 30)),
             sid))
        mysql.connection.commit()
        cur.close()
        flash('Student updated.', 'success')
        return redirect(url_for('student_list'))
    return render_template('student_form.html', student=student, announcements=get_announcements())

@app.route('/admin/students/delete/<sid>', methods=['POST'])
@admin_required
def delete_student(sid):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM students WHERE id = %s", (sid,))
    mysql.connection.commit()
    cur.close()
    flash('Student deleted.', 'success')
    return redirect(url_for('student_list'))

@app.route('/admin/students/reset-sessions', methods=['POST'])
@admin_required
def reset_sessions():
    cur = mysql.connection.cursor()
    cur.execute("UPDATE students SET session = 30")
    mysql.connection.commit()
    cur.close()
    flash('All sessions reset to 30.', 'success')
    return redirect(url_for('student_list'))

@app.route('/admin/sitin')
@admin_required
def current_sitin():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM sitin_records WHERE status = 'Active'")
    active = cur.fetchall()
    cur.close()
    return render_template('sitin_records.html', records=active, announcements=get_announcements())

@app.route('/admin/sitin/form', methods=['GET', 'POST'])
@admin_required
def sitin_form():
    sid = request.args.get('id', '')
    student = find_student(sid)
    if request.method == 'POST':
        sid2 = request.form.get('id_number', '').strip()
        student2 = find_student(sid2)
        if not student2:
            flash('Student not found.', 'danger')
            return redirect(url_for('sitin_form'))
        sit_id = next_sit_id()
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO sitin_records (sit_id, student_id, name, purpose, lab, session, status, date) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (sit_id, sid2,
             f"{student2['first']} {student2['last']}",
             request.form.get('purpose','').strip(),
             request.form.get('lab','').strip(),
             student2['session'],
             'Active',
             datetime.now().strftime("%Y-%m-%d %H:%M")))
        cur.execute("UPDATE students SET session = GREATEST(0, session - 1) WHERE id = %s", (sid2,))
        mysql.connection.commit()
        cur.close()
        flash(f"{student2['first']} {student2['last']} checked in successfully!", 'success')
        return redirect(url_for('current_sitin'))
    return render_template('sitin_form.html', student=student, announcements=get_announcements())

@app.route('/admin/sitin/checkout/<sit_id>', methods=['POST'])
@admin_required
def checkout(sit_id):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE sitin_records SET status = 'Done' WHERE sit_id = %s", (sit_id,))
    # Award 1 point to student when checked out
    cur.execute("""
        UPDATE students SET points = points + 1
        WHERE id = (SELECT student_id FROM sitin_records WHERE sit_id = %s)
    """, (sit_id,))
    mysql.connection.commit()
    cur.close()
    flash('Student checked out.', 'success')
    return redirect(url_for('current_sitin'))

@app.route('/admin/sitin/approve/<sit_id>', methods=['POST'])
@admin_required
def approve_sitin(sit_id):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE sitin_records SET status='Active' WHERE sit_id=%s", (sit_id,))
    mysql.connection.commit()
    cur.close()
    flash('Sit-in approved!', 'success')
    return redirect(url_for('lab_dashboard'))

@app.route('/admin/sitin/reject/<sit_id>', methods=['POST'])
@admin_required
def reject_sitin(sit_id):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE sitin_records SET status='Rejected' WHERE sit_id=%s", (sit_id,))
    mysql.connection.commit()
    cur.close()
    flash('Sit-in rejected.', 'danger')
    return redirect(url_for('lab_dashboard'))

@app.route('/admin/search')
@admin_required
def search_student():
    q = request.args.get('q', '').lower()
    results = []
    if q:
        cur = mysql.connection.cursor()
        cur.execute("SELECT * FROM students WHERE LOWER(id) LIKE %s OR LOWER(CONCAT(first,' ',last)) LIKE %s",
            (f'%{q}%', f'%{q}%'))
        results = cur.fetchall()
        cur.close()
    return render_template('search.html', results=results, query=q, announcements=get_announcements())

# ─── Reports routes ────────────────────────────────────────────────────────────

@app.route('/admin/reports')
@admin_required
def sitin_reports():
    status    = request.args.get('status', '')
    purpose   = request.args.get('purpose', '')
    date_from = request.args.get('date_from', '')
    date_to   = request.args.get('date_to', '')

    cur = mysql.connection.cursor()

    query = "SELECT * FROM sitin_records WHERE 1=1"
    params = []
    if status:
        query += " AND status = %s"; params.append(status)
    if purpose:
        query += " AND purpose = %s"; params.append(purpose)
    if date_from:
        query += " AND date >= %s"; params.append(date_from)
    if date_to:
        query += " AND date <= %s"; params.append(date_to + ' 23:59')
    query += " ORDER BY date DESC"
    cur.execute(query, params)
    sitin_records = cur.fetchall()

    cur.execute("SELECT * FROM students ORDER BY last")
    students = cur.fetchall()
    cur.close()

    return render_template('reports.html', sitin_records=sitin_records, students=students)

@app.route('/admin/reports/export/sitin/csv')
@admin_required
def export_sitin_csv():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM sitin_records ORDER BY date DESC")
    records = cur.fetchall()
    cur.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['Sit ID','Student ID','Name','Purpose','Lab','Session','Status','Date'])
    for r in records:
        writer.writerow([r['sit_id'], r['student_id'], r['name'], r['purpose'],
                         r['lab'], r['session'], r['status'], r['date']])

    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=sitin_records.csv'
    response.headers['Content-Type'] = 'text/csv'
    return response

@app.route('/admin/reports/export/students/csv')
@admin_required
def export_students_csv():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM students ORDER BY last")
    students = cur.fetchall()
    cur.close()

    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(['ID','Last Name','First Name','Middle Name','Year','Course','Email','Address','Sessions','Points'])
    for s in students:
        writer.writerow([s['id'], s['last'], s['first'], s['middle'], s['year'],
                         s['course'], s['email'], s['address'], s['session'],
                         s.get('points', 0)])

    response = make_response(output.getvalue())
    response.headers['Content-Disposition'] = 'attachment; filename=students_list.csv'
    response.headers['Content-Type'] = 'text/csv'
    return response

@app.route('/admin/reports/export/sitin/pdf')
@admin_required
def export_sitin_pdf():
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM sitin_records ORDER BY date DESC")
    records = cur.fetchall()
    cur.close()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("CCS Sit-in Records Report", styles['Title']))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 12))

    data = [['Sit ID','Student ID','Name','Purpose','Lab','Status','Date']]
    for r in records:
        data.append([r['sit_id'], r['student_id'], r['name'], r['purpose'],
                     r['lab'], r['status'], r['date']])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a237e')),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,0), (-1,0), 9),
        ('FONTSIZE',   (0,1), (-1,-1), 8),
        ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#e0e0e0')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f5f7ff')]),
        ('PADDING',    (0,0), (-1,-1), 5),
    ]))
    elements.append(table)
    doc.build(elements)

    buffer.seek(0)
    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = 'attachment; filename=sitin_records.pdf'
    response.headers['Content-Type'] = 'application/pdf'
    return response

@app.route('/admin/reports/export/students/pdf')
@admin_required
def export_students_pdf():
    from reportlab.lib.pagesizes import landscape, letter
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM students ORDER BY last")
    students = cur.fetchall()
    cur.close()

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=landscape(letter))
    styles = getSampleStyleSheet()
    elements = []

    elements.append(Paragraph("CCS Students List Report", styles['Title']))
    elements.append(Paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}", styles['Normal']))
    elements.append(Spacer(1, 12))

    data = [['ID','Last Name','First Name','Course','Year','Email','Sessions','Points']]
    for s in students:
        data.append([s['id'], s['last'], s['first'], s['course'], s['year'],
                     s['email'], s['session'], s.get('points', 0)])

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1a237e')),
        ('TEXTCOLOR',  (0,0), (-1,0), colors.white),
        ('FONTNAME',   (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE',   (0,0), (-1,0), 9),
        ('FONTSIZE',   (0,1), (-1,-1), 8),
        ('GRID',       (0,0), (-1,-1), 0.5, colors.HexColor('#e0e0e0')),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#f5f7ff')]),
        ('PADDING',    (0,0), (-1,-1), 5),
    ]))
    elements.append(table)
    doc.build(elements)

    buffer.seek(0)
    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = 'attachment; filename=students_list.pdf'
    response.headers['Content-Type'] = 'application/pdf'
    return response

@app.route('/admin/reports/export/sitin/docx')
@admin_required
def export_sitin_docx():
    from docx import Document

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM sitin_records ORDER BY date DESC")
    records = cur.fetchall()
    cur.close()

    doc = Document()
    doc.add_heading('CCS Sit-in Records Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['Sit ID','Student ID','Name','Purpose','Lab','Status','Date']):
        hdr[i].text = h

    for r in records:
        row = table.add_row().cells
        row[0].text = str(r['sit_id'])
        row[1].text = str(r['student_id'])
        row[2].text = str(r['name'])
        row[3].text = str(r['purpose'])
        row[4].text = str(r['lab'])
        row[5].text = str(r['status'])
        row[6].text = str(r['date'])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = 'attachment; filename=sitin_records.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response

@app.route('/admin/reports/export/students/docx')
@admin_required
def export_students_docx():
    from docx import Document

    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM students ORDER BY last")
    students = cur.fetchall()
    cur.close()

    doc = Document()
    doc.add_heading('CCS Students List Report', 0)
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    for i, h in enumerate(['ID','Last Name','First Name','Course','Year','Email','Sessions','Points']):
        hdr[i].text = h

    for s in students:
        row = table.add_row().cells
        row[0].text = str(s['id'])
        row[1].text = str(s['last'])
        row[2].text = str(s['first'])
        row[3].text = str(s['course'])
        row[4].text = str(s['year'])
        row[5].text = str(s['email'])
        row[6].text = str(s['session'])
        row[7].text = str(s.get('points', 0))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = make_response(buffer.read())
    response.headers['Content-Disposition'] = 'attachment; filename=students_list.docx'
    response.headers['Content-Type'] = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    return response

# ─── Analytics route ───────────────────────────────────────────────────────────

@app.route('/admin/analytics')
@admin_required
def analytics():
    cur = mysql.connection.cursor()

    cur.execute("SELECT COUNT(*) as cnt FROM students")
    total_students = cur.fetchone()['cnt']

    cur.execute("SELECT COUNT(*) as cnt FROM sitin_records")
    total_sitin = cur.fetchone()['cnt']

    cur.execute("SELECT COUNT(*) as cnt FROM sitin_records WHERE status='Active'")
    active_sitin = cur.fetchone()['cnt']

    cur.execute("SELECT COUNT(*) as cnt FROM reservations")
    total_reservations = cur.fetchone()['cnt']

    cur.execute("SELECT * FROM sitin_records")
    all_records = cur.fetchall()

    purpose_stats = {}
    for r in all_records:
        purpose_stats[r['purpose']] = purpose_stats.get(r['purpose'], 0) + 1

    lab_stats = {}
    for r in all_records:
        lab_stats[r['lab']] = lab_stats.get(r['lab'], 0) + 1

    status_stats = {}
    for r in all_records:
        status_stats[r['status']] = status_stats.get(r['status'], 0) + 1

    daily_stats = {}
    for i in range(6, -1, -1):
        day = (datetime.now() - timedelta(days=i)).strftime('%m/%d')
        daily_stats[day] = 0
    for r in all_records:
        try:
            day = datetime.strptime(r['date'][:10], '%Y-%m-%d').strftime('%m/%d')
            if day in daily_stats:
                daily_stats[day] += 1
        except:
            pass

    cur.execute("""
        SELECT s.*, COUNT(sr.sit_id) as sitin_count
        FROM students s
        LEFT JOIN sitin_records sr ON s.id = sr.student_id AND sr.status = 'Done'
        GROUP BY s.id
        ORDER BY s.points DESC, sitin_count DESC
        LIMIT 10
    """)
    top_students = cur.fetchall()
    cur.close()

    return render_template('analytics.html',
        total_students=total_students,
        total_sitin=total_sitin,
        active_sitin=active_sitin,
        total_reservations=total_reservations,
        purpose_stats=purpose_stats,
        lab_stats=lab_stats,
        status_stats=status_stats,
        daily_stats=daily_stats,
        top_students=top_students)

# ─── Rewards routes ────────────────────────────────────────────────────────────

@app.route('/admin/rewards')
@admin_required
def rewards():
    cur = mysql.connection.cursor()
    cur.execute("""
        SELECT s.*,
            (SELECT COUNT(*) FROM sitin_records WHERE student_id=s.id AND status='Done') as sitin_count,
            (SELECT COUNT(*) FROM reservations WHERE student_id=s.id AND status='Approved') as reservation_count
        FROM students s
        ORDER BY s.points DESC
    """)
    leaderboard = cur.fetchall()
    cur.close()
    return render_template('rewards.html', leaderboard=leaderboard)

@app.route('/admin/rewards/give', methods=['POST'])
@admin_required
def give_points():
    sid    = request.form.get('student_id', '').strip()
    pts    = int(request.form.get('points', 0))
    reason = request.form.get('reason', '').strip()
    student = find_student(sid)
    if not student:
        flash(f'Student {sid} not found.', 'danger')
    elif pts < 1:
        flash('Points must be at least 1.', 'danger')
    else:
        cur = mysql.connection.cursor()
        cur.execute("UPDATE students SET points = points + %s WHERE id = %s", (pts, sid))
        mysql.connection.commit()
        cur.close()
        flash(f'✅ Gave {pts} point(s) to {student["first"]} {student["last"]} for: {reason}', 'success')
    return redirect(url_for('rewards'))

# ─── Lab dashboard ─────────────────────────────────────────────────────────────

@app.route('/admin/lab')
@admin_required
def lab_dashboard():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM sitin_records WHERE status='Active'")
    active = cur.fetchall()
    cur.execute("SELECT * FROM sitin_records WHERE status='Pending'")
    pending = cur.fetchall()
    cur.execute("SELECT * FROM sitin_records ORDER BY date DESC LIMIT 20")
    logs = cur.fetchall()
    cur.close()

    return render_template('lab_dashboard.html',
        active=active,
        pending=pending,
        logs=logs,
        announcements=get_announcements())

@app.route('/admin/feedback')
@admin_required
def feedback_reports():
    return render_template('feedback.html', announcements=get_announcements())

# ─── Reservation routes ────────────────────────────────────────────────────────

@app.route('/admin/reservation')
@admin_required
def reservation():
    cur = mysql.connection.cursor()
    cur.execute("SELECT * FROM reservations ORDER BY id DESC")
    reservations = cur.fetchall()
    cur.close()
    return render_template('reservation.html', reservations=reservations, announcements=get_announcements())

@app.route('/admin/reservation/add', methods=['POST'])
@admin_required
def add_reservation():
    sid = request.form.get('student_id', '').strip()
    student = find_student(sid)
    name = f"{student['first']} {student['last']}" if student else "Unknown"
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO reservations (student_id, name, lab, purpose, date, time, status) VALUES (%s,%s,%s,%s,%s,%s,'Pending')",
        (sid, name, request.form.get('lab'), request.form.get('purpose'),
         request.form.get('date'), request.form.get('time')))
    mysql.connection.commit()
    cur.close()
    flash('Reservation added!', 'success')
    return redirect(url_for('reservation'))

@app.route('/admin/reservation/approve/<int:rid>', methods=['POST'])
@admin_required
def approve_reservation(rid):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE reservations SET status='Approved' WHERE id=%s", (rid,))
    # Award 2 points to student when reservation is approved
    cur.execute("""
        UPDATE students SET points = points + 2
        WHERE id = (SELECT student_id FROM reservations WHERE id = %s)
    """, (rid,))
    mysql.connection.commit()
    cur.close()
    flash('Reservation approved!', 'success')
    return redirect(url_for('reservation'))

@app.route('/admin/reservation/reject/<int:rid>', methods=['POST'])
@admin_required
def reject_reservation(rid):
    cur = mysql.connection.cursor()
    cur.execute("UPDATE reservations SET status='Rejected' WHERE id=%s", (rid,))
    mysql.connection.commit()
    cur.close()
    flash('Reservation rejected.', 'success')
    return redirect(url_for('reservation'))

@app.route('/admin/reservation/delete/<int:rid>', methods=['POST'])
@admin_required
def delete_reservation(rid):
    cur = mysql.connection.cursor()
    cur.execute("DELETE FROM reservations WHERE id=%s", (rid,))
    mysql.connection.commit()
    cur.close()
    flash('Reservation deleted.', 'success')
    return redirect(url_for('reservation'))

# ─── Student routes ────────────────────────────────────────────────────────────

@app.route('/student')
@student_required
def student_dashboard():
    student = find_student(session['user'])
    cur = mysql.connection.cursor()

    cur.execute("SELECT * FROM sitin_records WHERE student_id = %s", (session['user'],))
    my_records = cur.fetchall()

    cur.execute("SELECT * FROM reservations WHERE student_id = %s ORDER BY id DESC", (session['user'],))
    my_reservations = cur.fetchall()
    cur.close()

    notifications = []
    for r in my_records:
        if r['status'] == 'Active':
            notifications.append({
                'type': 'approved',
                'title': 'Sit-in Approved',
                'message': f"Your sit-in at {r['lab']} has been approved.",
                'date': r['date'],
                'is_read': False
            })
        elif r['status'] == 'Done':
            notifications.append({
                'type': 'done',
                'title': 'Sit-in Completed',
                'message': f"Your sit-in session at {r['lab']} is done.",
                'date': r['date'],
                'is_read': True
            })
    for r in my_reservations:
        if r['status'] == 'Approved':
            notifications.append({
                'type': 'approved',
                'title': 'Reservation Approved',
                'message': f"Your reservation for {r['lab']} on {r['date']} at {r['time']} is approved!",
                'date': r['date'],
                'is_read': False
            })
        elif r['status'] == 'Rejected':
            notifications.append({
                'type': 'rejected',
                'title': 'Reservation Rejected',
                'message': f"Your reservation for {r['lab']} on {r['date']} was rejected.",
                'date': r['date'],
                'is_read': False
            })

    return render_template('student_dashboard.html',
        student=student,
        announcements=get_announcements(),
        records=my_records,
        my_reservations=my_reservations,
        notifications=notifications)

@app.route('/student/sitin', methods=['POST'])
@student_required
def student_sitin_request():
    student = find_student(session['user'])
    if student and student['session'] > 0:
        sit_id = next_sit_id()
        cur = mysql.connection.cursor()
        cur.execute("INSERT INTO sitin_records (sit_id, student_id, name, purpose, lab, session, status, date) VALUES (%s,%s,%s,%s,%s,%s,%s,%s)",
            (sit_id, student['id'],
             f"{student['first']} {student['last']}",
             request.form.get('purpose','').strip(),
             request.form.get('lab','').strip(),
             student['session'],
             'Pending',
             datetime.now().strftime("%Y-%m-%d %H:%M")))
        mysql.connection.commit()
        cur.close()
        flash('Sit-in request submitted! Waiting for admin approval.', 'success')
    else:
        flash('No remaining sessions.', 'danger')
    return redirect(url_for('student_dashboard'))

@app.route('/student/reservation/add', methods=['POST'])
@student_required
def student_add_reservation():
    student = find_student(session['user'])
    name = f"{student['first']} {student['last']}"
    cur = mysql.connection.cursor()
    cur.execute("INSERT INTO reservations (student_id, name, lab, purpose, date, time, status) VALUES (%s,%s,%s,%s,%s,%s,'Pending')",
        (student['id'], name,
         request.form.get('lab'),
         request.form.get('purpose'),
         request.form.get('date'),
         request.form.get('time')))
    mysql.connection.commit()
    cur.close()
    flash('Reservation submitted! Waiting for admin approval.', 'success')
    return redirect(url_for('student_dashboard'))

if __name__ == '__main__':
    app.run(debug=True, port=5000)